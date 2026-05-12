#!/usr/bin/env python3
"""
Script para generar documento DOCX para entrega en Blackboard.

Requisitos:
    pip install python-docx

Uso:
    python generate_blackboard_submission.py
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os


def create_blackboard_submission():
    """Crea el documento DOCX para Blackboard."""
    
    doc = Document()
    
    # ==================== ENCABEZADO ====================
    title = doc.add_heading('🎤 Generative StandUp AI', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Aplicación de IA para generar rutinas de stand-up personalizadas')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.font.size = Pt(12)
    subtitle_format.font.italic = True
    
    doc.add_paragraph()  # Espacio
    
    # ==================== INFORMACIÓN DE ENTREGA ====================
    doc.add_heading('📋 Información de Entrega', 1)
    
    info_table = doc.add_table(rows=3, cols=2)
    info_table.style = 'Light Grid Accent 1'
    
    # Fila 1: GitHub
    info_table.rows[0].cells[0].text = '🔗 Repositorio GitHub'
    github_cell = info_table.rows[0].cells[1]
    github_para = github_cell.paragraphs[0]
    github_run = github_para.add_run('https://github.com/techboycr/StanUpNator')
    github_run.font.color.rgb = RGBColor(0, 0, 255)
    github_run.underline = True
    github_para.text = ''
    github_para.add_run('https://github.com/techboycr/StanUpNator')
    
    # Fila 2: App Desplegada
    info_table.rows[1].cells[0].text = '🌐 App Desplegada (Streamlit Cloud)'
    app_cell = info_table.rows[1].cells[1]
    app_para = app_cell.paragraphs[0]
    app_para.text = ''
    app_run = app_para.add_run('https://standup-ai-demo.streamlit.app/')
    app_run.font.color.rgb = RGBColor(0, 0, 255)
    app_run.underline = True
    
    # Fila 3: Fecha
    info_table.rows[2].cells[0].text = '📅 Fecha de Entrega'
    info_table.rows[2].cells[1].text = datetime.now().strftime('%d de %B de %Y')
    
    doc.add_paragraph()
    
    # ==================== DESCRIPCIÓN DEL PROYECTO ====================
    doc.add_heading('📖 Descripción del Proyecto', 1)
    
    desc_intro = doc.add_paragraph(
        'StandUp AI es una aplicación web interactiva que utiliza inteligencia artificial (Gemini) '
        'para crear rutinas de stand-up personalizadas. El sistema guía al usuario a través de un '
        'wizard de 3 pasos que combina:'
    )
    
    desc_points = [
        'Ingesta de contenido de YouTube (transcripciones como referencia)',
        'Entrevista guiada de 15 preguntas para extraer perfil y gustos del usuario',
        'Pipeline de generación multi-etapa (análisis → escritura → control de calidad)',
        'Revisión conversacional del perfil con chat en tiempo real',
        'Exportación en TXT y PDF'
    ]
    
    for point in desc_points:
        doc.add_paragraph(point, style='List Bullet')
    
    doc.add_paragraph()
    
    # ==================== TECNOLOGÍAS UTILIZADAS ====================
    doc.add_heading('💻 Tecnologías Utilizadas', 1)
    
    tech_table = doc.add_table(rows=7, cols=2)
    tech_table.style = 'Light Grid Accent 1'
    
    tech_data = [
        ('Frontend', 'Streamlit 1.35.0+'),
        ('LLM & Embeddings', 'Google Gemini 2.5 Flash'),
        ('Orquestación IA', 'LangChain 0.3.0+'),
        ('Vector Search & RAG', 'FAISS + Gemini Embeddings'),
        ('Integración YouTube', 'youtube-transcript-api, yt-dlp'),
        ('Generación PDF', 'fpdf2 2.7.9'),
        ('Testing', 'pytest 8.0.0, Playwright')
    ]
    
    for i, (category, tech) in enumerate(tech_data):
        tech_table.rows[i].cells[0].text = category
        tech_table.rows[i].cells[1].text = tech
    
    doc.add_paragraph()
    
    # ==================== FLUJO DE LA APLICACIÓN ====================
    doc.add_heading('🎯 Flujo de la Aplicación (Wizard)', 1)
    
    doc.add_heading('Paso 1: Ingesta de Videos', 2)
    doc.add_paragraph(
        'El usuario ingresa links de YouTube (uno por línea). El sistema valida, '
        'extrae transcripciones y construye un índice FAISS para recuperación de contexto.'
    )
    
    doc.add_heading('Paso 2: Entrevista Guiada', 2)
    doc.add_paragraph(
        'Chat interactivo con 15 preguntas sobre el estilo de comedia, temas favoritos, '
        'experiencias personales y público objetivo. Las respuestas se procesan para extraer '
        'un perfil estructurado del usuario.'
    )
    
    doc.add_heading('Paso 3: Perfil + Generación', 2)
    doc.add_paragraph(
        'El usuario revisa su perfil en un chat interactivo. Puede solicitar cambios que el '
        'sistema aplica en tiempo real. Cuando confirma con "sí", se ejecuta el pipeline de '
        'generación (análisis → rutina → QA) y puede descargar el resultado en TXT o PDF.'
    )
    
    doc.add_paragraph()
    
    # ==================== CARACTERÍSTICAS TÉCNICAS ====================
    doc.add_heading('⚙️ Características Técnicas', 1)
    
    features = [
        'Wizard secuencial de 3 pasos con barra de progreso',
        'Integración con Google Gemini API para análisis, generación y embeddings',
        'RAG (Retrieval-Augmented Generation) con FAISS y Gemini embeddings',
        'Extracción de transcripciones de YouTube con prioridad español/inglés',
        'Validación de duración de videos (ideal ≤30min, tolerancia ≤45min)',
        'Pipeline de generación con 3 agentes IA especializados:',
        '  • Psicólogo: Análisis de perfil',
        '  • Comediante: Escritura de rutina',
        '  • Productor: Control de calidad con métricas',
        'Fallbacks inteligentes cuando servicios remotos no están disponibles',
        'Interfaz de chat para revisión interactiva del perfil',
        'Exportación automática en TXT y PDF con sanitización de caracteres',
        'Mode producción/testing con control de variables de entorno'
    ]
    
    for feature in features:
        if feature.startswith('  •'):
            doc.add_paragraph(feature[3:], style='List Bullet 2')
        elif feature.endswith(':'):
            doc.add_paragraph(feature)
        else:
            doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_paragraph()
    
    # ==================== CAPTURAS DE PANTALLA ====================
    doc.add_heading('📸 Capturas de Pantalla', 1)
    
    doc.add_heading('Wizard - Paso 1 (Ingesta de Videos)', 2)
    doc.add_paragraph(
        'El usuario ingresa links de YouTube. El sistema valida, extrae transcripciones '
        'e inicializa el índice RAG para recuperación de contexto.'
    )
    
    screenshot_path_1 = 'docs/screenshots/wizard-step1.png'
    if os.path.exists(screenshot_path_1):
        try:
            doc.add_picture(screenshot_path_1, width=Inches(5.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f'[Screenshot no disponible: {e}]', style='Intense Quote')
    
    doc.add_paragraph()
    
    doc.add_heading('Wizard - Encabezado con Progreso', 2)
    doc.add_paragraph(
        'Barra de progreso visual mostrando el avance: Paso 1/3, 2/3 ó 3/3.'
    )
    
    screenshot_path_2 = 'docs/screenshots/wizard-header.png'
    if os.path.exists(screenshot_path_2):
        try:
            doc.add_picture(screenshot_path_2, width=Inches(5.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f'[Screenshot no disponible: {e}]', style='Intense Quote')
    
    doc.add_paragraph()
    
    # ==================== CÓDIGO RELEVANTE ====================
    doc.add_heading('💾 Código Relevante', 1)
    
    doc.add_heading('1. Pipeline de Generación Multi-Etapa', 2)
    pipeline_code = '''async def run_pipeline(user_profile, rag_retriever, selected_chat_model):
    """Ejecuta: análisis → rutina → control de calidad."""
    
    # Etapa 1: Análisis de perfil (temp=0.3)
    analysis = await analyze_user_profile(
        user_profile=user_profile,
        selected_model=selected_chat_model
    )
    
    # Etapa 2: Generación de rutina (temp=0.9)
    routine = await generate_routine(
        user_profile=user_profile,
        analysis=analysis,
        context=rag_retriever.invoke(user_profile.get('topics', '')),
        selected_model=selected_chat_model
    )
    
    # Etapa 3: QA (temp=0.1)
    qa_result = await quality_check(
        routine=routine,
        analysis=analysis,
        selected_model=selected_chat_model
    )
    
    return {
        'analysis': analysis,
        'routine': routine,
        'qa': qa_result
    }
'''
    doc.add_paragraph(pipeline_code, style='Normal')
    
    doc.add_heading('2. Interfaz de Paso 3 - Revisión de Perfil', 2)
    ui_code = '''def _render_step_3():
    """Revisión conversacional de perfil + generación automática."""
    
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.subheader("💬 Revisa tu perfil")
        user_message = st.chat_input("¿Cambios? O responde 'sí' para generar...")
        
    with col2:
        st.subheader("📋 Tu Perfil")
        profile_text = st.text_area(
            "Reporte", 
            value=st.session_state.profile_report,
            height=400
        )
    
    # Cuando usuario confirma
    if user_message and "sí" in user_message.lower():
        with st.spinner("Generando rutina..."):
            result = asyncio.run(
                run_pipeline(
                    st.session_state.user_profile,
                    st.session_state.rag_retriever,
                    selected_chat_model
                )
            )
        
        st.success("✅ ¡Rutina generada!")
        st.markdown(result['routine'])
        
        # Descargas
        st.download_button("📥 TXT", result['routine'])
        pdf_bytes = _build_pdf_bytes(result['routine'])
        st.download_button("📄 PDF", pdf_bytes)
'''
    doc.add_paragraph(ui_code, style='Normal')
    
    doc.add_heading('3. RAG con FAISS', 2)
    rag_code = '''def initialize_vectorstore(embedding_model):
    """Inicializa FAISS con ejemplos base."""
    
    base_samples = [
        "Set-up: La gente mira el teléfono en todas partes... "
        "Punchline: ¡Hasta aquí, mientras yo intento ser cómico!",
        # ... 4 ejemplos más de distintos estilos
    ]
    
    vectorstore = FAISS.from_texts(
        texts=base_samples,
        embedding=embedding_model
    )
    return vectorstore

def add_transcript_records(vectorstore, transcripts):
    """Indexa transcripciones en FAISS."""
    for transcript in transcripts:
        chunks = split_text(transcript, chunk_size=700, overlap=80)
        vectorstore.add_texts(chunks)
'''
    doc.add_paragraph(rag_code, style='Normal')
    
    doc.add_paragraph()
    
    # ==================== REQUISITOS DE INSTALACIÓN ====================
    doc.add_heading('📦 Instalación y Ejecución Local', 1)
    
    doc.add_heading('1. Clonar Repositorio', 2)
    doc.add_paragraph('git clone https://github.com/techboycr/StanUpNator.git', style='Intense Quote')
    doc.add_paragraph('cd StandUp-AI', style='Intense Quote')
    
    doc.add_heading('2. Crear Entorno Virtual', 2)
    doc.add_paragraph('python -m venv .venv', style='Intense Quote')
    doc.add_heading('Windows (PowerShell):', 3)
    doc.add_paragraph('.\.venv\Scripts\Activate.ps1', style='Intense Quote')
    doc.add_heading('Linux/macOS:', 3)
    doc.add_paragraph('source .venv/bin/activate', style='Intense Quote')
    
    doc.add_heading('3. Instalar Dependencias', 2)
    doc.add_paragraph('pip install -r requirements.txt', style='Intense Quote')
    
    doc.add_heading('4. Configurar Variables de Entorno', 2)
    doc.add_paragraph('Crear archivo .env con:', style='Normal')
    doc.add_paragraph('GOOGLE_API_KEY=tu_api_key_aqui', style='Intense Quote')
    
    doc.add_heading('5. Ejecutar Aplicación', 2)
    doc.add_paragraph('streamlit run app.py', style='Intense Quote')
    doc.add_paragraph('Luego abrir: http://localhost:8501', style='Normal')
    
    doc.add_paragraph()
    
    # ==================== TESTING ====================
    doc.add_heading('✅ Testing', 1)
    
    doc.add_paragraph('El proyecto incluye 18 pruebas unitarias que validan:')
    test_items = [
        'Extracción correcta de preguntas de entrevista',
        'Procesamiento de respuestas de usuario',
        'Funcionamiento del vector store FAISS',
        'Generación de rutina con fallbacks',
        'Control de calidad con métricas',
        'Ingesta de transcripciones de YouTube'
    ]
    for test in test_items:
        doc.add_paragraph(test, style='List Bullet')
    
    doc.add_paragraph('\nEjecutar pruebas:')
    doc.add_paragraph('python -m pytest tests/test_standup.py -v', style='Intense Quote')
    
    doc.add_paragraph()
    
    # ==================== DESPLIEGUE EN STREAMLIT CLOUD ====================
    doc.add_heading('☁️ Despliegue en Streamlit Community Cloud', 1)
    
    deploy_steps = [
        ('Paso 1', 'Ir a https://share.streamlit.io/ y clickear "New app"'),
        ('Paso 2', 'Seleccionar repositorio: techboycr/StanUpNator'),
        ('Paso 3', 'Seleccionar rama: main'),
        ('Paso 4', 'Seleccionar archivo: app.py'),
        ('Paso 5', 'Ir a Settings → Secrets y agregar GOOGLE_API_KEY'),
        ('Paso 6', 'Clickear Deploy y esperar 2-5 minutos'),
    ]
    
    for step, description in deploy_steps:
        p = doc.add_paragraph(description, style='List Number')
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_paragraph()
    
    # ==================== VALIDACIÓN ====================
    doc.add_heading('⚠️ Checklist Antes de Entregar', 1)
    
    checklist = [
        '✅ Repositorio GitHub es público',
        '✅ Link al repo está en README.md',
        '✅ App está desplegada en Streamlit Cloud',
        '✅ GOOGLE_API_KEY está configurado en Secrets',
        '✅ App.py, requirements.txt y README.md están en raíz del repo',
        '✅ Screenshots están en docs/screenshots/',
        '✅ Ambos links (GitHub y App) son públicos y accesibles',
        '✅ Wizard funciona correctamente en los 3 pasos',
        '✅ Descargas TXT y PDF funcionan',
        '✅ Todos los tests pasan (pytest tests/test_standup.py)'
    ]
    
    for item in checklist:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph()
    
    # ==================== LIMITACIONES Y MEJORAS FUTURAS ====================
    doc.add_heading('📝 Limitaciones y Roadmap', 1)
    
    doc.add_heading('Limitaciones Conocidas', 2)
    doc.add_paragraph(
        'Dependencia de servicios externos (Google Gemini, YouTube)',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Transcripciones pueden fallar por restricciones del video o red',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Calidad final depende de respuestas en la entrevista',
        style='List Bullet'
    )
    
    doc.add_heading('Mejoras Futuras', 2)
    doc.add_paragraph(
        'Persistencia de sesiones por usuario (base de datos)',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Historial de versiones de rutina',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Controles UI para ajustar tono/estructura de rutina',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Exportación adicional (Markdown, Docx)',
        style='List Bullet'
    )
    
    doc.add_paragraph()
    
    # ==================== FOOTER ====================
    doc.add_paragraph('_' * 80)
    
    footer = doc.add_paragraph(
        f'Documento generado: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")} | '
        'Proyecto: StandUp AI | Tecnología: Streamlit + Google Gemini'
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_format = footer.runs[0]
    footer_format.font.size = Pt(9)
    footer_format.font.italic = True
    footer_format.font.color.rgb = RGBColor(128, 128, 128)
    
    # ==================== GUARDAR ====================
    output_path = 'Entrega_Blackboard_StandUp-AI.docx'
    doc.save(output_path)
    print(f'✅ Documento creado: {output_path}')
    print(f'📍 Ruta: {os.path.abspath(output_path)}')
    print('\n✨ Listo para entregar en Blackboard')


if __name__ == '__main__':
    try:
        create_blackboard_submission()
    except ImportError:
        print('❌ Error: python-docx no está instalado')
        print('\nInstalar con:')
        print('  pip install python-docx')
    except Exception as e:
        print(f'❌ Error: {e}')
        import traceback
        traceback.print_exc()
